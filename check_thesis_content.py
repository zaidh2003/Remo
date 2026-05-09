"""
REMO Thesis Content Checker
Analyzes thesis content and provides improvement suggestions

Requirements:
    pip install python-docx

Usage:
    python check_thesis_content.py your_thesis.docx
"""

from docx import Document
import sys
import re

class ThesisContentChecker:
    """Check thesis content for completeness and quality"""
    
    def __init__(self, filename):
        self.doc = Document(filename)
        self.issues = []
        self.suggestions = []
        self.stats = {}
    
    def check_structure(self):
        """Check if all required sections are present"""
        required_sections = [
            'abstract', 'introduction', 'literature review', 'methodology',
            'implementation', 'results', 'discussion', 'conclusion', 'references'
        ]
        
        found_sections = []
        all_text = '\n'.join([p.text.lower() for p in self.doc.paragraphs])
        
        for section in required_sections:
            if section in all_text:
                found_sections.append(section)
            else:
                self.issues.append(f"Missing section: {section.title()}")
        
        return found_sections
    
    def count_words(self):
        """Count total words in document"""
        total_words = 0
        for paragraph in self.doc.paragraphs:
            words = paragraph.text.split()
            total_words += len(words)
        
        self.stats['total_words'] = total_words
        
        if total_words < 8000:
            self.issues.append(f"Word count too low: {total_words} (minimum 8000 recommended)")
        elif total_words > 15000:
            self.suggestions.append(f"Word count high: {total_words} (consider condensing)")
        
        return total_words
    
    def check_references(self):
        """Check reference section"""
        in_references = False
        ref_count = 0
        
        for paragraph in self.doc.paragraphs:
            if 'reference' in paragraph.text.lower() and paragraph.style.name.startswith('Heading'):
                in_references = True
                continue
            
            if in_references and paragraph.text.strip():
                ref_count += 1
        
        self.stats['references'] = ref_count
        
        if ref_count < 15:
            self.issues.append(f"Too few references: {ref_count} (minimum 15-20 recommended)")
        
        return ref_count
    
    def check_figures_tables(self):
        """Count figures and tables"""
        figure_count = 0
        table_count = len(self.doc.tables)
        
        for paragraph in self.doc.paragraphs:
            if 'figure' in paragraph.text.lower() and ':' in paragraph.text:
                figure_count += 1
        
        self.stats['figures'] = figure_count
        self.stats['tables'] = table_count
        
        if figure_count == 0:
            self.suggestions.append("Consider adding figures/diagrams to illustrate concepts")
        
        if table_count == 0:
            self.suggestions.append("Consider adding tables to present data")
        
        return figure_count, table_count
    
    def check_citations(self):
        """Check for in-text citations"""
        citation_patterns = [
            r'\([A-Z][a-z]+,?\s+\d{4}\)',  # (Author, 2024)
            r'\([A-Z][a-z]+\s+et\s+al\.,?\s+\d{4}\)',  # (Author et al., 2024)
            r'\[\d+\]',  # [1]
        ]
        
        citation_count = 0
        for paragraph in self.doc.paragraphs:
            for pattern in citation_patterns:
                citations = re.findall(pattern, paragraph.text)
                citation_count += len(citations)
        
        self.stats['citations'] = citation_count
        
        if citation_count < 20:
            self.issues.append(f"Too few citations: {citation_count} (more citations needed)")
        
        return citation_count
    
    def check_technical_terms(self):
        """Check if key technical terms are defined"""
        key_terms = [
            'next.js', 'react', 'firebase', 'firestore', 'groq', 'llama',
            'api', 'authentication', 'authorization', 'real-time', 'ai'
        ]
        
        all_text = '\n'.join([p.text.lower() for p in self.doc.paragraphs])
        
        mentioned_terms = []
        for term in key_terms:
            if term in all_text:
                mentioned_terms.append(term)
        
        self.stats['technical_terms'] = len(mentioned_terms)
        
        if len(mentioned_terms) < 5:
            self.suggestions.append("Include more technical details about the implementation")
        
        return mentioned_terms
    
    def check_remo_features(self):
        """Check if REMO features are mentioned"""
        features = [
            'shortage alert', 'shift swap', 'taxi management', 'emergency',
            'schedule', 'inventory', 'forecast', 'multilingual', 'branch'
        ]
        
        all_text = '\n'.join([p.text.lower() for p in self.doc.paragraphs])
        
        mentioned_features = []
        for feature in features:
            if feature in all_text:
                mentioned_features.append(feature)
        
        missing_features = set(features) - set(mentioned_features)
        
        if missing_features:
            self.suggestions.append(f"Consider mentioning these features: {', '.join(missing_features)}")
        
        return mentioned_features
    
    def check_metrics(self):
        """Check if performance metrics are included"""
        metric_keywords = [
            'improvement', 'reduction', 'increase', 'faster', 'efficiency',
            'performance', 'metric', 'result', 'measurement'
        ]
        
        all_text = '\n'.join([p.text.lower() for p in self.doc.paragraphs])
        
        metrics_mentioned = sum(1 for keyword in metric_keywords if keyword in all_text)
        
        if metrics_mentioned < 3:
            self.issues.append("Include more quantitative results and performance metrics")
        
        return metrics_mentioned
    
    def generate_report(self):
        """Generate comprehensive report"""
        print("\n" + "=" * 70)
        print("REMO THESIS CONTENT ANALYSIS REPORT")
        print("=" * 70)
        
        # Statistics
        print("\n📊 DOCUMENT STATISTICS")
        print("-" * 70)
        print(f"Total Words:        {self.stats.get('total_words', 0):,}")
        print(f"References:         {self.stats.get('references', 0)}")
        print(f"In-text Citations:  {self.stats.get('citations', 0)}")
        print(f"Figures:            {self.stats.get('figures', 0)}")
        print(f"Tables:             {self.stats.get('tables', 0)}")
        print(f"Technical Terms:    {self.stats.get('technical_terms', 0)}")
        
        # Issues
        if self.issues:
            print("\n❌ ISSUES FOUND")
            print("-" * 70)
            for i, issue in enumerate(self.issues, 1):
                print(f"{i}. {issue}")
        else:
            print("\n✅ No critical issues found!")
        
        # Suggestions
        if self.suggestions:
            print("\n💡 SUGGESTIONS FOR IMPROVEMENT")
            print("-" * 70)
            for i, suggestion in enumerate(self.suggestions, 1):
                print(f"{i}. {suggestion}")
        
        # Overall Assessment
        print("\n📋 OVERALL ASSESSMENT")
        print("-" * 70)
        
        score = 100
        score -= len(self.issues) * 10
        score -= len(self.suggestions) * 5
        score = max(0, score)
        
        if score >= 80:
            status = "✅ GOOD - Ready for review"
        elif score >= 60:
            status = "⚠️  FAIR - Needs improvements"
        else:
            status = "❌ POOR - Significant work needed"
        
        print(f"Content Quality Score: {score}/100")
        print(f"Status: {status}")
        
        print("\n" + "=" * 70)
        print("RECOMMENDATIONS")
        print("=" * 70)
        print("""
1. Review all issues and address them
2. Implement suggested improvements
3. Have your supervisor review the content
4. Run the formatter: python format_thesis.py your_thesis.docx
5. Proofread for grammar and spelling
6. Check all citations and references
7. Verify all figures and tables are properly captioned
8. Ensure consistent terminology throughout
        """)


def main():
    if len(sys.argv) < 2:
        print("Usage: python check_thesis_content.py your_thesis.docx")
        sys.exit(1)
    
    filename = sys.argv[1]
    
    print(f"\nAnalyzing thesis: {filename}")
    print("Please wait...")
    
    try:
        checker = ThesisContentChecker(filename)
        
        print("\n1. Checking document structure...")
        checker.check_structure()
        
        print("2. Counting words...")
        checker.count_words()
        
        print("3. Analyzing references...")
        checker.check_references()
        
        print("4. Counting figures and tables...")
        checker.check_figures_tables()
        
        print("5. Checking citations...")
        checker.check_citations()
        
        print("6. Verifying technical content...")
        checker.check_technical_terms()
        
        print("7. Checking REMO features...")
        checker.check_remo_features()
        
        print("8. Analyzing metrics...")
        checker.check_metrics()
        
        checker.generate_report()
        
    except Exception as e:
        print(f"\n✗ Error analyzing thesis: {str(e)}")
        sys.exit(1)


if __name__ == "__main__":
    main()
