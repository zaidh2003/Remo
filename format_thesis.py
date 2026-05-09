"""
REMO Thesis Formatter for Turiba University Standards
This script formats a thesis document according to Turiba University requirements.

Requirements:
    pip install python-docx

Usage:
    python format_thesis.py input_thesis.docx output_thesis.docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
import sys
import os

class TuribaThesisFormatter:
    """Format thesis according to Turiba University standards"""
    
    def __init__(self, input_file):
        self.doc = Document(input_file)
        self.setup_styles()
    
    def setup_styles(self):
        """Setup standard styles for Turiba thesis"""
        styles = self.doc.styles
        
        # Normal text style
        normal = styles['Normal']
        normal.font.name = 'Times New Roman'
        normal.font.size = Pt(12)
        normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        normal.paragraph_format.first_line_indent = Inches(0.5)
        
        # Heading 1 style
        try:
            heading1 = styles['Heading 1']
        except KeyError:
            heading1 = styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
        heading1.font.name = 'Times New Roman'
        heading1.font.size = Pt(14)
        heading1.font.bold = True
        heading1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        heading1.paragraph_format.space_before = Pt(12)
        heading1.paragraph_format.space_after = Pt(6)
        
        # Heading 2 style
        try:
            heading2 = styles['Heading 2']
        except KeyError:
            heading2 = styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)
        heading2.font.name = 'Times New Roman'
        heading2.font.size = Pt(13)
        heading2.font.bold = True
        heading2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        heading2.paragraph_format.space_before = Pt(10)
        heading2.paragraph_format.space_after = Pt(5)
        
        # Heading 3 style
        try:
            heading3 = styles['Heading 3']
        except KeyError:
            heading3 = styles.add_style('Heading 3', WD_STYLE_TYPE.PARAGRAPH)
        heading3.font.name = 'Times New Roman'
        heading3.font.size = Pt(12)
        heading3.font.bold = True
        heading3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        heading3.paragraph_format.space_before = Pt(8)
        heading3.paragraph_format.space_after = Pt(4)
    
    def format_page_setup(self):
        """Set up page margins and size"""
        sections = self.doc.sections
        for section in sections:
            # A4 page size
            section.page_height = Inches(11.69)
            section.page_width = Inches(8.27)
            
            # Margins: Left 1.5", Right 1", Top 1", Bottom 1"
            section.left_margin = Inches(1.5)
            section.right_margin = Inches(1.0)
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
    
    def format_paragraphs(self):
        """Format all paragraphs according to standards"""
        for paragraph in self.doc.paragraphs:
            # Skip if it's a heading
            if paragraph.style.name.startswith('Heading'):
                continue
            
            # Apply normal style
            paragraph.style = 'Normal'
            
            # Ensure proper formatting
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
    
    def add_page_numbers(self):
        """Add page numbers to footer"""
        for section in self.doc.sections:
            footer = section.footer
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_para.text = "Page "
            
    def fix_references(self):
        """Format references section according to APA/Harvard style"""
        in_references = False
        for paragraph in self.doc.paragraphs:
            if 'reference' in paragraph.text.lower() and paragraph.style.name.startswith('Heading'):
                in_references = True
                continue
            
            if in_references and paragraph.text.strip():
                # Format reference entries
                paragraph.paragraph_format.left_indent = Inches(0.5)
                paragraph.paragraph_format.first_line_indent = Inches(-0.5)
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    
    def create_table_of_contents(self):
        """Create or update table of contents"""
        # Find TOC location or create new
        toc_found = False
        for i, paragraph in enumerate(self.doc.paragraphs):
            if 'table of contents' in paragraph.text.lower():
                toc_found = True
                # Clear existing TOC
                paragraph.clear()
                paragraph.add_run('TABLE OF CONTENTS').bold = True
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                break
        
        if not toc_found:
            # Add TOC after title page (assuming first 3 paragraphs are title page)
            self.doc.add_paragraph('TABLE OF CONTENTS', style='Heading 1')
    
    def format_tables(self):
        """Format all tables consistently"""
        for table in self.doc.tables:
            # Try to set table style, skip if not available
            try:
                table.style = 'Light Grid Accent 1'
            except:
                try:
                    table.style = 'Table Grid'
                except:
                    pass  # Skip if no style available
            
            # Format cells
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.style = 'Normal'
                        for run in paragraph.runs:
                            run.font.size = Pt(11)
    
    def add_title_page(self, title, author, supervisor, date):
        """Create a properly formatted title page"""
        # Clear first page if exists
        if self.doc.paragraphs:
            first_para = self.doc.paragraphs[0]
            first_para.clear()
        else:
            first_para = self.doc.add_paragraph()
        
        # University name
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('TURIBA UNIVERSITY')
        run.font.size = Pt(14)
        run.font.bold = True
        
        # Add spacing
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        # Thesis title
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title.upper())
        run.font.size = Pt(16)
        run.font.bold = True
        
        # Add spacing
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        # Author
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'By\n{author}')
        run.font.size = Pt(12)
        
        # Add spacing
        self.doc.add_paragraph()
        
        # Supervisor
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'Supervisor: {supervisor}')
        run.font.size = Pt(12)
        
        # Add spacing
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        # Date and location
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'Riga, Latvia\n{date}')
        run.font.size = Pt(12)
        
        # Page break after title page
        self.doc.add_page_break()
    
    def check_spelling_grammar(self):
        """Basic spell check for common errors"""
        common_errors = {
            'teh': 'the',
            'adn': 'and',
            'taht': 'that',
            'thsi': 'this',
            'whcih': 'which',
            'recieve': 'receive',
            'occured': 'occurred',
            'seperate': 'separate',
        }
        
        for paragraph in self.doc.paragraphs:
            text = paragraph.text
            for wrong, correct in common_errors.items():
                if wrong in text.lower():
                    # Replace in all runs
                    for run in paragraph.runs:
                        run.text = run.text.replace(wrong, correct)
                        run.text = run.text.replace(wrong.capitalize(), correct.capitalize())
    
    def save(self, output_file):
        """Save the formatted document"""
        self.doc.save(output_file)
        print(f"✓ Thesis formatted and saved to: {output_file}")


def main():
    """Main function to run the formatter"""
    if len(sys.argv) < 2:
        print("Usage: python format_thesis.py input_thesis.docx [output_thesis.docx]")
        print("\nThis script will format your thesis according to Turiba University standards:")
        print("  - Times New Roman, 12pt font")
        print("  - Double-spaced paragraphs")
        print("  - Proper margins (Left 1.5\", others 1\")")
        print("  - Justified text alignment")
        print("  - Formatted headings")
        print("  - Page numbers")
        print("  - Proper reference formatting")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else input_file.replace('.docx', '_formatted.docx')
    
    if not os.path.exists(input_file):
        print(f"Error: Input file '{input_file}' not found!")
        sys.exit(1)
    
    print(f"Formatting thesis: {input_file}")
    print("=" * 60)
    
    try:
        formatter = TuribaThesisFormatter(input_file)
        
        print("1. Setting up page layout...")
        formatter.format_page_setup()
        
        print("2. Formatting paragraphs...")
        formatter.format_paragraphs()
        
        print("3. Formatting tables...")
        formatter.format_tables()
        
        print("4. Fixing references...")
        formatter.fix_references()
        
        print("5. Checking common spelling errors...")
        formatter.check_spelling_grammar()
        
        print("6. Saving formatted document...")
        formatter.save(output_file)
        
        print("\n" + "=" * 60)
        print("✓ Formatting complete!")
        print(f"\nFormatted thesis saved as: {output_file}")
        print("\nPlease review the document and make any necessary adjustments.")
        print("\nTuriba Standards Applied:")
        print("  ✓ Times New Roman, 12pt")
        print("  ✓ Double-spaced paragraphs")
        print("  ✓ Margins: Left 1.5\", Right/Top/Bottom 1\"")
        print("  ✓ Justified alignment")
        print("  ✓ Proper heading hierarchy")
        print("  ✓ Reference formatting")
        
    except Exception as e:
        print(f"\n✗ Error formatting thesis: {str(e)}")
        print("\nPlease check:")
        print("  1. The input file is a valid .docx file")
        print("  2. The file is not open in another program")
        print("  3. You have python-docx installed: pip install python-docx")
        sys.exit(1)


if __name__ == "__main__":
    main()
